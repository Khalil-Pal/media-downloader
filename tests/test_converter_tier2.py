from __future__ import annotations

import asyncio
import json
import os
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from urllib.parse import unquote, urlparse
from unittest.mock import AsyncMock, patch

os.environ.setdefault("BOT_TOKEN", "123456:test-token")
os.environ.setdefault("PYTHON_DOTENV_DISABLED", "1")

import fitz
import pymupdf4llm
from docx import Document
from openpyxl import load_workbook
from pptx import Presentation

from config.settings import settings
from handlers import convert_handler
from handlers.convert_handler import handle_convertible_file
from services import db
from services.converter import (
    ConversionError,
    _convert_with_libreoffice,
    _is_heading_1,
    _normalized_table_rows,
    cleanup_conversion_session,
    convert_file,
    create_conversion_session,
    friendly_conversion_error,
    supported_targets,
    tier2_caveat_key,
)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PROJECT_ROOT = Path(__file__).resolve().parent.parent


def _make_text_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    lines = [
        "Sandy Squirrel Tier 2 conversion test",
        "This PDF contains a normal extractable text layer.",
        "The page deliberately contains no table.",
    ]
    for index, text in enumerate(lines):
        page.insert_text((72, 72 + index * 24), text, fontsize=12)
    document.save(path)
    document.close()


def _make_table_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 48), "Sandy Squirrel usage table", fontsize=14)
    rows = [
        ("Item", "Count", "Status"),
        ("Videos", "10", "Ready"),
        ("Audio", "7", "Ready"),
    ]
    for row_index, row in enumerate(rows):
        y = 90 + row_index * 26
        for x, value in zip((72, 220, 360), row):
            page.insert_text((x, y), value, fontsize=12)
    document.save(path)
    document.close()


def _make_image_only_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    pixmap = fitz.Pixmap(fitz.csRGB, (0, 0, 300, 120), 0)
    pixmap.clear_with(230)
    page.insert_image(fitz.Rect(72, 72, 372, 192), pixmap=pixmap)
    document.save(path)
    document.close()


def _make_heading_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Downloads", level=1)
    document.add_paragraph("Download supported media")
    document.add_paragraph("Extract audio", style="List Bullet")
    document.add_heading("Conversions", level=1)
    document.add_paragraph("Convert documents and images")
    document.save(path)


def _make_plain_docx(path: Path) -> None:
    document = Document()
    for index in range(12):
        document.add_paragraph(f"Paragraph {index + 1}")
    document.save(path)


class Tier2StaticTests(unittest.TestCase):
    def test_supported_targets_extend_existing_tier1_options(self) -> None:
        self.assertEqual(
            [option.target_format for option in supported_targets("sample.pdf", PDF_MIME)],
            ["docx", "pptx", "md", "xlsx"],
        )
        self.assertEqual(
            [option.target_format for option in supported_targets("sample.docx", DOCX_MIME)],
            ["pdf", "pptx"],
        )
        self.assertEqual(
            [option.target_format for option in supported_targets("sample.png", "image/png")],
            ["jpg", "webp"],
        )

    def test_caveats_are_limited_to_tier2(self) -> None:
        self.assertEqual(
            tier2_caveat_key("sample.pdf", "docx", PDF_MIME),
            "conversion_tier2_caveat",
        )
        self.assertEqual(
            tier2_caveat_key("sample.docx", "pptx", DOCX_MIME),
            "conversion_tier2_caveat_docx_pptx",
        )
        self.assertIsNone(tier2_caveat_key("sample.docx", "pdf", DOCX_MIME))
        self.assertIsNone(tier2_caveat_key("sample.png", "jpg", "image/png"))

    def test_table_normalization_drops_titles_and_rejects_noise(self) -> None:
        self.assertIsNone(_normalized_table_rows([["Title"], ["Body"]]))
        self.assertIsNone(_normalized_table_rows([["A", ""], ["", ""]]))
        self.assertEqual(
            _normalized_table_rows(
                [
                    ["Sandy Squirrel usage table", "", ""],
                    ["Type", "Count", "Status"],
                    ["Video", "10", "Ready"],
                ]
            ),
            [
                ["Type", "Count", "Status"],
                ["Video", "10", "Ready"],
            ],
        )
        self.assertEqual(
            _normalized_table_rows(
                [[float("nan"), "Name", "Count"], [None, "Video", "10"]]
            ),
            [["Name", "Count"], ["Video", "10"]],
        )

    def test_heading_detection_uses_stable_style_id(self) -> None:
        paragraph = SimpleNamespace(
            style=SimpleNamespace(style_id="Heading1", name="Localized heading name")
        )
        self.assertTrue(_is_heading_1(paragraph))

    def test_new_error_codes_use_specific_translation_keys(self) -> None:
        self.assertEqual(
            friendly_conversion_error(ConversionError("scanned_pdf")),
            "conversion_error_pdf_scanned",
        )
        self.assertEqual(
            friendly_conversion_error(ConversionError("no_tables")),
            "conversion_error_pdf_no_tables",
        )

    def test_tier2_translations_exist_in_every_locale(self) -> None:
        required_keys = {
            "conversion_tier2_caveat",
            "conversion_tier2_caveat_docx_pptx",
            "conversion_error_pdf_scanned",
            "conversion_error_pdf_no_tables",
            "conversion_error_pdf_text_failed",
        }
        for language in ("en", "ar", "ru"):
            with self.subTest(language=language):
                translations = json.loads(
                    (PROJECT_ROOT / f"{language}.json").read_text(encoding="utf-8")
                )
                self.assertFalse(required_keys - translations.keys())


class Tier2ConversionTests(unittest.IsolatedAsyncioTestCase):
    async def asyncSetUp(self) -> None:
        self._temp_dir = tempfile.TemporaryDirectory()
        self._original_download_path = settings.download_path
        object.__setattr__(settings, "download_path", Path(self._temp_dir.name))

        db._pool = None
        db._memory_users.clear()
        db._memory_user_plans.clear()
        db._memory_payments.clear()
        db._memory_cancelled_payment_refs.clear()

    async def asyncTearDown(self) -> None:
        object.__setattr__(settings, "download_path", self._original_download_path)
        self._temp_dir.cleanup()

    def _assert_download_root_empty(self) -> None:
        self.assertEqual(list(settings.download_path.iterdir()), [])

    async def test_pdf_to_markdown_real_conversion_and_cleanup(self) -> None:
        session = create_conversion_session()
        try:
            source = session / "text.pdf"
            _make_text_pdf(source)

            output = await convert_file(source, "md", PDF_MIME)

            self.assertGreater(output.stat().st_size, 0)
            self.assertIn("Sandy Squirrel", output.read_text(encoding="utf-8"))
        finally:
            cleanup_conversion_session(session)
        self._assert_download_root_empty()

    async def test_pdf_to_markdown_falls_back_to_plain_text(self) -> None:
        session = create_conversion_session()
        try:
            source = session / "fallback.pdf"
            _make_text_pdf(source)
            with self.assertLogs("services.converter", level="WARNING"):
                with patch.object(
                    pymupdf4llm,
                    "to_markdown",
                    side_effect=RuntimeError("primary extractor failed"),
                ):
                    output = await convert_file(source, "md", PDF_MIME)

            text = output.read_text(encoding="utf-8")
            self.assertIn("## Page 1", text)
            self.assertIn("Sandy Squirrel", text)
        finally:
            cleanup_conversion_session(session)
        self._assert_download_root_empty()

    async def test_pdf_to_excel_real_table_and_zero_table_rejection(self) -> None:
        session = create_conversion_session()
        try:
            source = session / "table.pdf"
            _make_table_pdf(source)
            output = await convert_file(source, "xlsx", PDF_MIME)

            workbook = load_workbook(output, read_only=True)
            try:
                self.assertEqual(workbook.sheetnames, ["Table 1"])
                values = list(workbook["Table 1"].values)
                self.assertEqual(values[0], ("Item", "Count", "Status"))
            finally:
                workbook.close()
        finally:
            cleanup_conversion_session(session)
        self._assert_download_root_empty()

        session = create_conversion_session()
        try:
            source = session / "no-table.pdf"
            _make_text_pdf(source)
            with self.assertRaisesRegex(ConversionError, "^no_tables$"):
                await convert_file(source, "xlsx", PDF_MIME)
        finally:
            cleanup_conversion_session(session)
        self._assert_download_root_empty()

    async def test_image_only_pdf_is_rejected_and_cleaned(self) -> None:
        session = create_conversion_session()
        try:
            source = session / "scan.pdf"
            _make_image_only_pdf(source)
            with patch(
                "services.converter._convert_with_libreoffice",
                new=AsyncMock(),
            ) as libreoffice:
                with self.assertRaisesRegex(ConversionError, "^scanned_pdf$"):
                    await convert_file(source, "docx", PDF_MIME)
            libreoffice.assert_not_awaited()
        finally:
            cleanup_conversion_session(session)
        self._assert_download_root_empty()

    async def test_docx_to_pptx_heading_and_chunk_fallback(self) -> None:
        for source_name, maker, expected_slides in (
            ("headings.docx", _make_heading_docx, 2),
            ("plain.docx", _make_plain_docx, 3),
        ):
            with self.subTest(source_name=source_name):
                session = create_conversion_session()
                try:
                    source = session / source_name
                    maker(source)
                    output = await convert_file(source, "pptx", DOCX_MIME)
                    presentation = Presentation(output)
                    self.assertEqual(len(presentation.slides), expected_slides)
                    self.assertGreater(output.stat().st_size, 0)
                finally:
                    cleanup_conversion_session(session)
                self._assert_download_root_empty()

    async def test_libreoffice_profiles_are_unique_and_removed(self) -> None:
        sessions = [create_conversion_session(), create_conversion_session()]
        profile_paths: list[Path] = []
        commands: list[list[str]] = []

        async def fake_run_process(command, _tool_name, timeout):
            commands.append(command)
            profile_arg = next(
                item for item in command if item.startswith("-env:UserInstallation=")
            )
            profile_uri = urlparse(profile_arg.split("=", 1)[1])
            profile_path = unquote(profile_uri.path)
            if os.name == "nt" and profile_path.startswith("/") and ":" in profile_path[:4]:
                profile_path = profile_path[1:]
            profile_paths.append(Path(profile_path))
            target_format = command[command.index("--convert-to") + 1].split(":", 1)[0]
            source = Path(command[-1])
            source.with_suffix("." + target_format).write_bytes(b"converted")
            await asyncio.sleep(0)

        try:
            sources = []
            for index, session in enumerate(sessions):
                source = session / f"sample-{index}.pdf"
                _make_text_pdf(source)
                sources.append(source)

            with patch("services.converter._run_process", side_effect=fake_run_process):
                outputs = await asyncio.gather(
                    *(
                        _convert_with_libreoffice(
                            source,
                            "docx",
                            180,
                            "writer_pdf_import",
                            "Office Open XML Text",
                        )
                        for source in sources
                    )
                )

            self.assertEqual(len(set(profile_paths)), 2)
            self.assertTrue(all(output.stat().st_size > 0 for output in outputs))
            self.assertFalse(any(path.exists() for path in profile_paths))
            self.assertTrue(
                all("--infilter=writer_pdf_import" in command for command in commands)
            )
            self.assertTrue(
                all(
                    "docx:Office Open XML Text" in command
                    for command in commands
                )
            )
        finally:
            for session in sessions:
                cleanup_conversion_session(session)
        self._assert_download_root_empty()

    async def test_downloader_mode_gates_pdf_and_docx(self) -> None:
        user_id = 701
        await db.set_user_lang(user_id, "en")
        await db.set_user_mode(user_id, "downloader")

        for file_name, mime_type in (
            ("sample.pdf", PDF_MIME),
            ("sample.docx", DOCX_MIME),
        ):
            with self.subTest(file_name=file_name):
                message = SimpleNamespace(
                    from_user=SimpleNamespace(id=user_id),
                    document=SimpleNamespace(
                        file_id="file-id",
                        file_name=file_name,
                        mime_type=mime_type,
                        file_size=100,
                    ),
                    video=None,
                    audio=None,
                    answer=AsyncMock(),
                )

                await handle_convertible_file(message)

                self.assertIn("Converter mode", message.answer.await_args.args[0])

    async def test_tier2_callback_uses_shared_rate_limiter(self) -> None:
        user_id = 702
        await db.set_user_lang(user_id, "en")
        request = convert_handler.ConversionRequest(
            user_id=user_id,
            file_id="file-id",
            file_name="sample.pdf",
            mime_type=PDF_MIME,
            file_size=100,
            targets=("md",),
        )
        token = convert_handler._store_request(request)
        callback = SimpleNamespace(
            from_user=SimpleNamespace(id=user_id),
            data=f"convert:md:{token}",
            message=SimpleNamespace(),
            answer=AsyncMock(),
        )

        try:
            with patch.object(
                convert_handler.rate_limiter,
                "check",
                new=AsyncMock(return_value=(False, "rate limited")),
            ) as check:
                await convert_handler.cb_convert(callback, SimpleNamespace())

            check.assert_awaited_once_with(user_id)
            callback.answer.assert_awaited_once_with("rate limited", show_alert=True)
        finally:
            convert_handler._conversion_store.pop(token, None)


if __name__ == "__main__":
    unittest.main()
